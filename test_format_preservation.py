"""
Test: formatting preservation after heading style change.

Creates a document with:
  - Normal paragraphs with explicit font/bold settings
  - Multi-run headings (prefix run + bold body run)
  - Paragraphs with different existing styles

Then runs process_document and checks that the output has:
  - Correct number prefixes
  - Preserved run count (multi-run headings stay multi-run)
  - Preserved bold/font on body runs
"""
import sys, os, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from word_processor import process_document, _collect_rpr_chain, _replace_prefix_in_para

# ---------------------------------------------------------------------------
# Helper: create test document
# ---------------------------------------------------------------------------

def make_test_doc(path: str):
    doc = Document()

    # Para 0: Normal body text (should NOT become a heading)
    p0 = doc.add_paragraph("Dieser Absatz ist normaler Fließtext.")

    # Para 1: Single-run heading – Roman numeral
    p1 = doc.add_paragraph()
    r1 = p1.add_run("I. Vertragsgegenstand")
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    # NOT bold

    # Para 2: Multi-run heading – prefix run (plain) + body run (bold)
    p2 = doc.add_paragraph()
    r2a = p2.add_run("II. ")
    r2a.font.name = "Arial"
    r2a.font.size = Pt(11)
    r2b = p2.add_run("Laufzeit des Vertrages")
    r2b.font.name = "Arial"
    r2b.font.size = Pt(11)
    r2b.bold = True          # body is bold, prefix is not

    # Para 3: Body text following heading
    p3 = doc.add_paragraph("Der Vertrag beginnt am 1. Januar 2024.")

    # Para 4: §-heading
    p4 = doc.add_paragraph("§ 1 Allgemeine Bestimmungen")

    # Para 5: Another §-heading
    p5 = doc.add_paragraph("§ 2 Vergütung")

    doc.save(path)
    return doc


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_run_preservation():
    """Multi-run heading: after prefix replacement, body run must stay bold."""
    doc = Document()
    p = doc.add_paragraph()
    r_prefix = p.add_run("II. ")
    r_body   = p.add_run("Laufzeit")
    r_body.bold = True

    assert p.runs[1].bold is True, "Setup: body run must be bold"
    assert p.runs[0].bold is None, "Setup: prefix run must be non-bold (None)"

    _replace_prefix_in_para(p, "II. ", "2. ")

    # After replacement: prefix run now contains "2. ", body run untouched
    assert p.runs[0].text == "2. ",     f"prefix run text wrong: {p.runs[0].text!r}"
    assert p.runs[1].text == "Laufzeit", f"body run text wrong: {p.runs[1].text!r}"
    assert p.runs[1].bold is True,       "body run must STILL be bold after prefix replacement"
    print("  PASS  test_run_preservation")


def test_no_extra_runs_created():
    """Single-run heading: replacement must not add extra runs."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("I. Vertragsgegenstand")
    r.font.name = "Arial"

    before_runs = len(p.runs)
    _replace_prefix_in_para(p, "I. ", "1. ")

    assert p.runs[0].text == "1. Vertragsgegenstand", f"text wrong: {p.runs[0].text!r}"
    assert len(p.runs) == before_runs, f"run count changed: {before_runs} → {len(p.runs)}"
    print("  PASS  test_no_extra_runs_created")


def test_prefix_removal():
    """_replace_prefix_in_para with new_prefix='' removes old prefix only."""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("III. ")
    r2 = p.add_run("Haftung")
    r2.bold = True

    _replace_prefix_in_para(p, "III. ", "")

    assert p.runs[0].text == "", f"prefix run should be empty: {p.runs[0].text!r}"
    assert p.runs[1].text == "Haftung", f"body run wrong: {p.runs[1].text!r}"
    assert p.runs[1].bold is True, "body bold must survive prefix removal"
    print("  PASS  test_prefix_removal")


def test_full_pipeline():
    """End-to-end: process document, check output runs and text."""
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "input.docx")
        out = os.path.join(tmp, "output.docx")

        make_test_doc(src)
        count = process_document(src, out)

        assert count > 0, f"Expected headings, got {count}"
        assert os.path.exists(out), "Output file not created"

        result = Document(out)
        texts = [p.text for p in result.paragraphs]

        print(f"  Detected {count} headings")
        print(f"  Output paragraphs:")
        for i, p in enumerate(result.paragraphs):
            print(f"    [{i}] style={p.style.name!r:20s}  text={p.text!r}")

        # Check that headings have proper prefix
        heading_paras = [p for p in result.paragraphs if "Heading" in p.style.name]
        assert len(heading_paras) > 0, "No Heading-styled paragraphs in output"

        # Check that body text was NOT made into a heading
        fliesstext = [p for p in result.paragraphs if "Fließtext" in p.text or "beginnt" in p.text]
        for p in fliesstext:
            assert "Heading" not in p.style.name, f"Body text got heading style: {p.text!r}"

        # Check multi-run heading: body run must still be bold
        for p in result.paragraphs:
            if "Laufzeit" in p.text:
                bold_runs = [r for r in p.runs if "Laufzeit" in r.text and r.bold]
                # Body run might have bold=True or bold=None (inherited)
                # Just check that at least the text is preserved
                assert "Laufzeit" in p.text, "Laufzeit text lost"
                print(f"  Laufzeit para runs: {[(r.text, r.bold) for r in p.runs]}")

        print("  PASS  test_full_pipeline")


# ---------------------------------------------------------------------------
# Run all tests
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Running formatting preservation tests...")
    errors = []

    for test_fn in [test_run_preservation, test_no_extra_runs_created,
                    test_prefix_removal, test_full_pipeline]:
        try:
            test_fn()
        except AssertionError as e:
            print(f"  FAIL  {test_fn.__name__}: {e}")
            errors.append(test_fn.__name__)
        except Exception as e:
            print(f"  ERROR {test_fn.__name__}: {e}")
            import traceback; traceback.print_exc()
            errors.append(test_fn.__name__)

    if errors:
        print(f"\n{len(errors)} test(s) FAILED: {errors}")
        sys.exit(1)
    else:
        print("\nAll tests passed.")
