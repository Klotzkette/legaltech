"""
Word Processor – Core logic for heading detection and standardisation.

Detects heading paragraphs (by style, text pattern, or formatting heuristic),
remaps them to a consistent 1. / 1.1 / 1.1.1 … numbering scheme using
Word's built-in multilevel list linked to Heading 1-9 styles, and optionally
records all style changes as OOXML tracked changes (w:pPrChange).
"""

import copy
import datetime
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Accepted file extensions
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".doc", ".docx", ".rtf", ".txt"}

# ---------------------------------------------------------------------------
# Heading style name → level mapping  (English + German)
# ---------------------------------------------------------------------------

_HEADING_NAME_MAP: Dict[str, int] = {}
for _i in range(1, 10):
    _HEADING_NAME_MAP[f"heading {_i}"] = _i
    _HEADING_NAME_MAP[f"überschrift {_i}"] = _i
    _HEADING_NAME_MAP[f"uberschrift {_i}"] = _i   # without umlaut
    _HEADING_NAME_MAP[f"header {_i}"] = _i
    _HEADING_NAME_MAP[f"title {_i}"] = _i
    _HEADING_NAME_MAP[f"kopfzeile {_i}"] = _i

# ---------------------------------------------------------------------------
# Numbering TYPE detection  (not level — level is assigned by context)
# ---------------------------------------------------------------------------

_ROMAN_RE = r'(?:XIV|XIII|XII|XI|IX|VIII|VII|VI|IV|III|II|X|I|V)'

_TYPE_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r'^\d+\.\d+\.\d+\.\d+[\s\)]'),  "DECIMAL_4"),
    (re.compile(r'^\d+\.\d+\.\d+[\s\)]'),         "DECIMAL_3"),
    (re.compile(r'^\d+\.\d+[\s\)]'),              "DECIMAL_2"),
    (re.compile(r'^\d{1,2}\.\s+\S'),              "DECIMAL_1"),
    (re.compile(rf'^{_ROMAN_RE}\.\s+\S'),          "ROMAN"),
    (re.compile(r'^[A-Z]{1,3}\.\s+\S'),            "CAPITAL"),
    (re.compile(r'^§\s*\d+'),                      "PARAGRAPH"),
    (re.compile(r'^Art\.?\s*\d+', re.I),           "ARTICLE"),
    (re.compile(r'^Ziff?\.?\s*\d+', re.I),         "ZIFFER"),
    (re.compile(r'^(?:Teil|Kapitel|Abschnitt|Titel)\s+(?:\d+|[IVX]+)\b', re.I), "CHAPTER"),
    # (1)/(a)/a)/aa) intentionally absent – German legal body text, NOT headings.
]

_TYPE_PRIORITY: Dict[str, int] = {
    "PARAGRAPH":    10,
    "CHAPTER":      10,
    "ROMAN":        20,
    "ARTICLE":      20,
    "CAPITAL":      30,
    "ZIFFER":       35,
    "DECIMAL_1":    35,
    "DECIMAL_2":    40,
    "DECIMAL_3":    50,
    "DECIMAL_4":    60,
    "BOLD_ONLY":    15,
    "WORD_AUTO":    25,
}

_STRIP_PATTERNS: List[re.Pattern] = [
    re.compile(r'^\d+\.\d+\.\d+\.\d+[\s\)]+'),
    re.compile(r'^\d+\.\d+\.\d+[\s\)]+'),
    re.compile(r'^\d+\.\d+[\s\)]+'),
    re.compile(r'^\d{1,2}\.\s+'),
    re.compile(rf'^{_ROMAN_RE}\.\s+', re.IGNORECASE),
    re.compile(r'^[A-Z]{1,3}\.\s+'),
    re.compile(r'^§\s*\d+\s*[:\-–]?\s*'),
    re.compile(r'^Art\.?\s*\d+\s*[:\-–]?\s*', re.I),
    re.compile(r'^Ziff?\.?\s*\d+\s*[:\-–]?\s*', re.I),
    re.compile(r'^(?:Teil|Kapitel|Abschnitt|Titel)\s+(?:\d+|[IVX]+)\s*[:\-–]?\s*', re.I),
]


# ---------------------------------------------------------------------------
# Low-level heading detection helpers
# ---------------------------------------------------------------------------

def _level_from_style(para) -> Optional[int]:
    name = para.style.name.lower().strip()
    if name in _HEADING_NAME_MAP:
        return _HEADING_NAME_MAP[name]
    m = re.search(r'(?:heading|überschrift|uberschrift|header)\s+(\d+)', name)
    if m:
        level = int(m.group(1))
        return level if 1 <= level <= 9 else None
    return None


def _style_font_size_pt(para) -> Optional[float]:
    try:
        sz = para.style.font.size
        if sz is not None:
            return sz.pt
    except Exception:
        pass
    return None


def _has_word_auto_numbering(para) -> bool:
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    numPr_el = pPr.find(qn("w:numPr"))
    if numPr_el is None:
        return False
    numId_el = numPr_el.find(qn("w:numId"))
    if numId_el is None:
        return False
    try:
        return int(numId_el.get(qn("w:val"), "0")) > 0
    except ValueError:
        return False


def _detect_numbering_type(text: str) -> Optional[str]:
    for pattern, type_name in _TYPE_PATTERNS:
        if pattern.match(text):
            return type_name
    return None


def _is_heading_heuristic(para) -> bool:
    """Conservative bold / ALL-CAPS / large-font heading detector.

    Strict thresholds avoid false positives in German legal body text.
    """
    text = para.text.strip()
    if not text:
        return False
    if text.startswith(("•", "–", "-", "*", "(", "[")):
        return False
    if re.search(r"[.,;!?:]\s*$", text):
        return False

    words = text.split()

    # Bold-only: ALL runs bold AND short AND few words
    if len(text) <= 80 and len(words) <= 7:
        style_bold = False
        try:
            style_bold = para.style.font.bold is True
        except Exception:
            pass

        def _run_is_bold(r) -> bool:
            if r.bold is True:
                return True
            if r.bold is False:
                return False
            return style_bold

        non_empty = [r for r in para.runs if r.text.strip()]
        if non_empty and all(_run_is_bold(r) for r in non_empty):
            return True

    # ALL-CAPS
    if len(text) <= 60 and len(words) <= 5 and text == text.upper() and any(
        c.isalpha() for c in text
    ):
        return True

    # Large font
    pt = _style_font_size_pt(para)
    if pt is not None and pt >= 14 and len(text) <= 120:
        return True

    return False


def _context_aware_levels(
    heading_info: List[Tuple[int, str]],
) -> Dict[int, int]:
    """Assign heading levels based on numbering-type hierarchy in the document."""
    if not heading_info:
        return {}

    present_types: List[str] = []
    first_para_idx: Dict[str, int] = {}
    for para_idx, ntype in heading_info:
        if ntype not in present_types:
            present_types.append(ntype)
            first_para_idx[ntype] = para_idx

    adjusted_prio: Dict[str, int] = {t: _TYPE_PRIORITY.get(t, 99)
                                      for t in present_types}

    for t in present_types:
        prio = adjusted_prio[t]
        first = first_para_idx[t]
        count_before = sum(
            1 for pi, nt in heading_info
            if nt != t and pi < first and _TYPE_PRIORITY.get(nt, 99) > prio
        )
        if count_before >= 3:
            first_type = present_types[0]
            adjusted_prio[t] = adjusted_prio[first_type]

    present_types.sort(key=lambda t: (adjusted_prio[t], first_para_idx[t]))

    type_to_level: Dict[str, int] = {}
    current_level = 0
    last_priority = -1
    for t in present_types:
        p = adjusted_prio[t]
        if p != last_priority:
            current_level += 1
            last_priority = p
        type_to_level[t] = current_level

    return {idx: type_to_level[ntype] for idx, ntype in heading_info}


def strip_manual_numbering(text: str) -> str:
    stripped, _ = _split_prefix(text)
    return stripped


def _split_prefix(text: str) -> Tuple[str, str]:
    for pattern in _STRIP_PATTERNS:
        m = pattern.match(text)
        if m:
            body = text[m.end():].lstrip()
            prefix = text[: len(text) - len(body)]
            return body, prefix
    return text, ""


# ---------------------------------------------------------------------------
# Document heading analysis
# ---------------------------------------------------------------------------

def detect_headings(doc: Document) -> Dict[int, int]:
    """Detect headings via styles, numbering patterns, and formatting heuristics."""
    style_headings: Dict[int, int] = {}
    type_headings: List[Tuple[int, str]] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_level = _level_from_style(para)
        if style_level is not None:
            style_headings[i] = style_level
            continue

        if len(text) <= 120:
            num_type = _detect_numbering_type(text)
            if num_type is not None:
                body_check, _ = _split_prefix(text)
                body_core = body_check.rstrip(".")
                is_sentence = (
                    body_check.endswith(".")
                    and len(body_check) > 15
                    and len(body_core.split()) >= 3
                )
                if not is_sentence:
                    type_headings.append((i, num_type))
                    continue

        if _has_word_auto_numbering(para) and len(text) <= 200:
            type_headings.append((i, "WORD_AUTO"))
            continue

        if _is_heading_heuristic(para):
            type_headings.append((i, "BOLD_ONLY"))

    type_levels = _context_aware_levels(type_headings)

    headings: Dict[int, int] = {}
    headings.update(style_headings)
    headings.update(type_levels)
    return headings


def normalize_levels(headings: Dict[int, int]) -> Dict[int, int]:
    if not headings:
        return headings
    unique = sorted(set(headings.values()))
    remap = {old: new + 1 for new, old in enumerate(unique)}
    return {i: remap[lvl] for i, lvl in headings.items()}


def compute_number_strings(headings: Dict[int, int]) -> Dict[int, str]:
    counters: List[int] = [0] * 10
    result: Dict[int, str] = {}
    for idx in sorted(headings.keys()):
        level = headings[idx]
        for parent in range(level - 1):
            if counters[parent] == 0:
                counters[parent] = 1
        counters[level - 1] += 1
        for deeper in range(level, 10):
            counters[deeper] = 0
        parts = [str(counters[i]) for i in range(level)]
        result[idx] = ".".join(parts) + ("." if level == 1 else "")
    return result


# ---------------------------------------------------------------------------
# OOXML numbering helpers  (kept for reference; not used in main pipeline)
# ---------------------------------------------------------------------------

def _build_abstract_num(abstract_num_id: int, style_ids: Dict[int, str]) -> object:
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    mt = OxmlElement("w:multiLevelType")
    mt.set(qn("w:val"), "multilevel")
    abstract_num.append(mt)

    for i in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(i))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")

        if i == 0:
            lvl_text_val = "%1."
        else:
            lvl_text_val = ".".join(f"%{j + 1}" for j in range(i + 1))

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), lvl_text_val)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")

        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str((i + 1) * 720))
        ind.set(qn("w:hanging"), "360")
        pPr.append(ind)

        sid = style_ids.get(i + 1)
        ordered = [start, num_fmt]
        if sid:
            pStyle_el = OxmlElement("w:pStyle")
            pStyle_el.set(qn("w:val"), sid)
            ordered.append(pStyle_el)
        ordered += [lvl_text, lvl_jc, pPr]
        for child in ordered:
            lvl.append(child)

        abstract_num.append(lvl)

    return abstract_num


def _ensure_numbering_part(doc: Document):
    try:
        return doc.part.numbering_part
    except Exception:
        pass

    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    numbering_element = parse_xml(f'<w:numbering xmlns:w="{_W_NS}"/>'.encode())
    ct = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.numbering+xml"
    )
    numbering_part = NumberingPart(
        PackURI("/word/numbering.xml"), ct, numbering_element, None
    )
    doc.part.relate_to(numbering_part, RT.NUMBERING)
    return numbering_part


def setup_numbering(doc: Document) -> int:
    np = _ensure_numbering_part(doc)
    num_el = np._element

    style_ids: Dict[int, str] = {}
    for lvl in range(1, 10):
        try:
            style_ids[lvl] = doc.styles[f"Heading {lvl}"].style_id
        except KeyError:
            pass

    abstract_nums = num_el.findall(qn("w:abstractNum"))
    max_abs = max(
        (int(a.get(qn("w:abstractNumId"), -1)) for a in abstract_nums),
        default=-1,
    )
    new_abs_id = max_abs + 1

    nums = num_el.findall(qn("w:num"))
    max_num = max(
        (int(n.get(qn("w:numId"), 0)) for n in nums),
        default=0,
    )
    new_num_id = max_num + 1

    abstract_num = _build_abstract_num(new_abs_id, style_ids)
    if nums:
        nums[0].addprevious(abstract_num)
    else:
        num_el.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abs_id_el = OxmlElement("w:abstractNumId")
    abs_id_el.set(qn("w:val"), str(new_abs_id))
    num.append(abs_id_el)
    num_el.append(num)

    return new_num_id


def link_styles_to_numbering(doc: Document, num_id: int) -> None:
    for level in range(1, 10):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue

        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.element.append(pPr)

        for existing in pPr.findall(qn("w:numPr")):
            pPr.remove(existing)

        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level - 1))
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(ilvl)
        numPr.append(numId_el)
        pPr.insert(0, numPr)


# ---------------------------------------------------------------------------
# Formatting freeze – XML-level (the only reliable approach)
# ---------------------------------------------------------------------------
#
# python-docx font properties (run.font.name, run.font.size, run.bold …)
# return None when the value is inherited from the style.  Walking the
# python-docx style chain often still returns None because style fonts use
# theme-font references (w:theme="minorLatinFont") that python-docx doesn't
# resolve to a real font name.
#
# The only reliable fix is to work directly with OOXML elements:
#   1. Collect the raw <w:rFonts>, <w:sz>, <w:b>, <w:i>, <w:color> elements
#      from the full style chain (run rPr → para mark rPr → style chain → docDefaults).
#   2. Copy them verbatim into each run's rPr BEFORE the style change.
# This preserves theme-font references, half-point sizes, etc. exactly as they
# were rendered.
#
# ---------------------------------------------------------------------------

# OOXML W namespace (used for direct element access below)
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP = "{" + _W + "}"            # element-level prefix, e.g. _WP + "rFonts"

# rPr child tags we consider visually significant
_RPR_VISUAL_TAGS = ("rFonts", "sz", "szCs", "b", "bCs", "i", "iCs",
                    "color", "lang", "kern", "spacing", "w", "vertAlign",
                    "strike", "dstrike", "u", "highlight")


def _collect_rpr_chain(para, run, doc) -> Dict[str, object]:
    """Collect effective rPr child elements for *run* in *para*.

    Walks the inheritance chain from most-specific to least-specific:
      run's own rPr  →  para-mark rPr (pPr/rPr)  →  style chain  →  docDefaults

    Returns {local_tag_name: lxml_element} – the first (most-specific) value
    found for each tag wins.
    """
    result: Dict[str, object] = {}

    def _scan_rpr(rPr_el):
        if rPr_el is None:
            return
        for child in rPr_el:
            # lxml tag looks like "{namespace}localname"
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local in _RPR_VISUAL_TAGS and local not in result:
                result[local] = child

    # 1. Run's own rPr
    run_rPr_el = run._r.find(f"{_WP}rPr")
    _scan_rpr(run_rPr_el)

    # 1b. Walk the run's character style (w:rStyle) chain, if any.
    #     Character styles can carry bold/color/font that must be preserved.
    if run_rPr_el is not None:
        rStyle_el = run_rPr_el.find(f"{_WP}rStyle")
        if rStyle_el is not None:
            rStyle_val = rStyle_el.get(f"{_WP}val")
            if rStyle_val:
                _visited_char: set = set()
                for sty_candidate in doc.styles:
                    try:
                        if sty_candidate.style_id == rStyle_val:
                            cs = sty_candidate
                            while cs is not None:
                                cid = getattr(cs, "style_id", id(cs))
                                if cid in _visited_char:
                                    break
                                _visited_char.add(cid)
                                try:
                                    _scan_rpr(cs.element.find(f"{_WP}rPr"))
                                except Exception:
                                    pass
                                cs = getattr(cs, "base_style", None)
                            break
                    except Exception:
                        pass

    # 2. Paragraph-mark rPr  (pPr/rPr)
    pPr = para._p.find(f"{_WP}pPr")
    if pPr is not None:
        _scan_rpr(pPr.find(f"{_WP}rPr"))

    # 3. Style chain (most-specific first: para.style → base_style → …)
    sty = para.style
    visited: set = set()
    while sty is not None:
        sid = getattr(sty, "style_id", id(sty))
        if sid in visited:
            break
        visited.add(sid)
        try:
            _scan_rpr(sty.element.find(f"{_WP}rPr"))
        except Exception:
            pass
        sty = getattr(sty, "base_style", None)

    # 4. Document defaults
    try:
        rPrDef = doc.element.find(f".//{_WP}rPrDefault")
        if rPrDef is not None:
            _scan_rpr(rPrDef.find(f"{_WP}rPr"))
    except Exception:
        pass

    return result


def _freeze_run_formatting(para, doc) -> None:
    """Pin all effective character formatting explicitly onto every run.

    Must be called BEFORE changing ``para.style``.  After this call every
    run has fully explicit rPr attributes, so the subsequent style change
    cannot alter the visual appearance (font family, size, bold, italic,
    colour, language …).

    Also calls _freeze_para_properties to block heading-style spacing/indent.
    """
    for run in para.runs:
        eff = _collect_rpr_chain(para, run, doc)

        # Ensure the run has an rPr container
        run_rPr = run._r.find(f"{_WP}rPr")
        if run_rPr is None:
            run_rPr = OxmlElement("w:rPr")
            run._r.insert(0, run_rPr)

        # If this run has a character style (w:rStyle), those style properties
        # are already resolved into `eff` via the rStyle chain walk above.
        # Do NOT add fallback overrides (b=0, i=0, color=auto) for runs with
        # a rStyle because the character style is the authoritative source.
        has_rStyle = run_rPr.find(f"{_WP}rStyle") is not None

        # Copy every effective property that is not already explicit on this run
        for tag in _RPR_VISUAL_TAGS:
            if run_rPr.find(f"{_WP}{tag}") is None:
                if tag in eff:
                    run_rPr.append(copy.deepcopy(eff[tag]))
                elif has_rStyle:
                    # Character style may provide this property; don't override
                    pass
                elif tag == "b":
                    # No bold defined anywhere in chain → explicitly NOT bold
                    # (Heading styles are bold by default; this blocks that)
                    b = OxmlElement("w:b")
                    b.set(qn("w:val"), "0")
                    run_rPr.append(b)
                elif tag == "i":
                    i = OxmlElement("w:i")
                    i.set(qn("w:val"), "0")
                    run_rPr.append(i)
                elif tag == "color":
                    # No explicit color anywhere → write auto to block blue
                    # heading theme color
                    c = OxmlElement("w:color")
                    c.set(qn("w:val"), "auto")
                    run_rPr.append(c)

    _freeze_para_properties(para)


def _freeze_para_properties(para) -> None:
    """Write effective paragraph-level spacing/indent/justification explicitly.

    Must be called BEFORE changing ``para.style``.  Explicit pPr children
    have higher priority than style defaults, so the heading style's spacing
    and indentation cannot override them after the change.
    """
    # Ensure pPr exists
    pPr = para._p.find(f"{_WP}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)

    for tag in ("ind", "spacing", "jc"):
        # If already explicit on this paragraph, it will survive the style
        # change automatically – nothing to do.
        if pPr.find(f"{_WP}{tag}") is not None:
            continue

        # Find effective value in style chain (most-specific first)
        eff_el = None
        sty = para.style
        visited: set = set()
        while sty is not None:
            sid = getattr(sty, "style_id", id(sty))
            if sid in visited:
                break
            visited.add(sid)
            try:
                sty_pPr = sty.element.find(f"{_WP}pPr")
                if sty_pPr is not None:
                    child = sty_pPr.find(f"{_WP}{tag}")
                    if child is not None:
                        eff_el = copy.deepcopy(child)
                        break
            except Exception:
                pass
            sty = getattr(sty, "base_style", None)

        if eff_el is not None:
            # For spacing: ensure both w:before and w:after are present so
            # the heading style cannot augment a partial element.
            if tag == "spacing":
                if eff_el.get(qn("w:before")) is None:
                    eff_el.set(qn("w:before"), "0")
                if eff_el.get(qn("w:after")) is None:
                    eff_el.set(qn("w:after"), "0")
            pPr.append(eff_el)
        elif tag == "ind":
            # No indent found → write zero to block heading indent
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:hanging"), "0")
            pPr.append(ind)
        elif tag == "spacing":
            # No spacing found → write zero to block heading space-before
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:before"), "0")
            sp.set(qn("w:after"), "0")
            pPr.append(sp)
        elif tag == "jc":
            # No justification found → write left to block any heading center/justify
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "left")
            pPr.append(jc)


# ---------------------------------------------------------------------------
# Surgical prefix replacement
# ---------------------------------------------------------------------------

def _replace_prefix_in_para(para, old_prefix: str, new_prefix: str) -> None:
    """Replace the numbering prefix while preserving ALL run formatting.

    Walks the runs consuming ``old_prefix`` character-by-character across run
    boundaries.  Only the run(s) that actually contain the old prefix are
    modified; every run that holds purely body text is left completely
    untouched (its bold, italic, colour, etc. are preserved exactly).
    """
    if not para.runs:
        if new_prefix:
            para.add_run(new_prefix)
        return
    if old_prefix == new_prefix:
        return
    if not old_prefix:
        para.runs[0].text = new_prefix + para.runs[0].text
        return

    remaining = old_prefix
    new_written = False
    for run in para.runs:
        if not remaining:
            break
        t = run.text
        if not t:
            continue

        if remaining.startswith(t):
            # This run is fully inside the old prefix
            remaining = remaining[len(t):]
            if not new_written:
                run.text = new_prefix
                new_written = True
            else:
                run.text = ""

        elif t.startswith(remaining):
            # The prefix ends inside this run; the rest is body text
            body_tail = t[len(remaining):]
            remaining = ""
            if not new_written:
                run.text = new_prefix + body_tail
                new_written = True
            else:
                run.text = body_tail

        else:
            # Mismatch (e.g. tracked-change runs interspersed) → safe fallback:
            # collapse to single run to guarantee correctness
            full = para.text
            body = full[len(old_prefix):] if full.startswith(old_prefix) else full
            para.runs[0].text = new_prefix + body
            for r in para.runs[1:]:
                r.text = ""
            return

    if not new_written and new_prefix:
        # Safety: prefix spanned more than all run text (shouldn't happen)
        para.runs[0].text = new_prefix + para.runs[0].text[len(old_prefix):]


# ---------------------------------------------------------------------------
# Apply heading styles (with optional Track Changes recording)
# ---------------------------------------------------------------------------

def apply_heading_styles(
    doc: Document,
    headings: Dict[int, int],
    track_changes: bool = False,
    author: str = "Word-Gliederungs-Retter",
) -> None:
    """Apply Heading 1-9 styles and write literal 1. / 1.1 / 1.1.1 … prefixes.

    Freeze-before-change guarantees the heading style switch cannot alter the
    visual appearance of any paragraph.  Only the numbering prefix changes.

    Track-changes mode records every modification as OOXML revision marks.
    """
    number_strings = compute_number_strings(headings)
    date_str = (
        datetime.datetime.now(datetime.timezone.utc)
        .strftime("%Y-%m-%dT%H:%M:%SZ")
    )

    existing_ids = [
        int(el.get(qn("w:id")))
        for el in doc.element.body.iter()
        if el.get(qn("w:id")) is not None and el.get(qn("w:id")).isdigit()
    ]
    change_id = max(existing_ids, default=0) + 1

    for para_idx, level in sorted(headings.items()):
        para = doc.paragraphs[para_idx]
        style_name = f"Heading {level}"

        try:
            target_style = doc.styles[style_name]
        except KeyError:
            continue

        # ── Compute new text ──────────────────────────────────────────────
        number_str = number_strings.get(para_idx, "")
        raw_text   = para.text
        body, old_prefix = _split_prefix(raw_text)
        if not body:
            body       = raw_text
            old_prefix = ""
        new_prefix = f"{number_str} " if number_str else ""
        new_text   = new_prefix + body

        # ── Determine whether the style actually needs to change ──────────
        current_style_id = None
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None:
                current_style_id = ps.get(qn("w:val"))
        style_already_correct = (current_style_id == target_style.style_id)

        # ── Freeze formatting ONLY when a style change is about to happen ─
        # If the style is already correct we skip the freeze to avoid
        # unnecessarily expanding every run's rPr with redundant attributes.
        if not style_already_correct:
            _freeze_run_formatting(para, doc)

        # ── Apply Heading style ────────────────────────────────────────────
        if track_changes and not style_already_correct:
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                para._p.insert(0, pPr)
            orig_pPr = copy.deepcopy(pPr)
            for ch in orig_pPr.findall(qn("w:pPrChange")):
                orig_pPr.remove(ch)
            ps = pPr.find(qn("w:pStyle"))
            if ps is None:
                ps = OxmlElement("w:pStyle")
                pPr.insert(0, ps)
            ps.set(qn("w:val"), target_style.style_id)
            for ch in pPr.findall(qn("w:pPrChange")):
                pPr.remove(ch)
            pPrChange = OxmlElement("w:pPrChange")
            pPrChange.set(qn("w:id"), str(change_id))
            pPrChange.set(qn("w:author"), author)
            pPrChange.set(qn("w:date"), date_str)
            pPrChange.append(orig_pPr)
            pPr.append(pPrChange)
            change_id += 1
        elif not style_already_correct:
            para.style = target_style

        # ── Rewrite paragraph text with the new number prefix ─────────────
        if raw_text == new_text:
            continue

        first_run_el = para._p.find(qn("w:r"))
        first_rPr = None
        if first_run_el is not None:
            first_rPr = first_run_el.find(qn("w:rPr"))

        if track_changes:
            if old_prefix:
                del_el = OxmlElement("w:del")
                del_el.set(qn("w:id"), str(change_id))
                del_el.set(qn("w:author"), author)
                del_el.set(qn("w:date"), date_str)
                del_run = OxmlElement("w:r")
                if first_rPr is not None:
                    del_run.append(copy.deepcopy(first_rPr))
                del_text = OxmlElement("w:delText")
                del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                del_text.text = old_prefix
                del_run.append(del_text)
                del_el.append(del_run)
                if first_run_el is not None:
                    first_run_el.addprevious(del_el)
                else:
                    para._p.append(del_el)
                change_id += 1

            if new_prefix:
                ins_el = OxmlElement("w:ins")
                ins_el.set(qn("w:id"), str(change_id))
                ins_el.set(qn("w:author"), author)
                ins_el.set(qn("w:date"), date_str)
                ins_run = OxmlElement("w:r")
                if first_rPr is not None:
                    ins_run.append(copy.deepcopy(first_rPr))
                ins_text = OxmlElement("w:t")
                ins_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                ins_text.text = new_prefix
                ins_run.append(ins_text)
                ins_el.append(ins_run)
                if first_run_el is not None:
                    first_run_el.addprevious(ins_el)
                else:
                    para._p.append(ins_el)
                change_id += 1

            # Remove old prefix; body runs are NOT touched
            _replace_prefix_in_para(para, old_prefix, "")
        else:
            # Direct mode: surgical in-place prefix replacement
            _replace_prefix_in_para(para, old_prefix, new_prefix)


# ---------------------------------------------------------------------------
# .doc / .rtf → .docx conversion via LibreOffice
# ---------------------------------------------------------------------------

def convert_doc_to_docx(doc_path: str) -> str:
    """Convert a .doc or .rtf file to .docx using LibreOffice (headless).

    The converted file is written to a temporary directory.  The caller is
    responsible for cleaning it up via ``cleanup_converted_tmp`` or by passing
    the returned path to ``_cleanup_tmp_docx`` after loading the document.
    """
    tmp_dir = tempfile.mkdtemp(prefix="gliederungsretter_")
    soffice_candidates = [
        "soffice",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for soffice in soffice_candidates:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", tmp_dir, doc_path],
                capture_output=True,
                timeout=60,
            )
            if result.returncode == 0:
                out = Path(tmp_dir) / (Path(doc_path).stem + ".docx")
                if out.exists():
                    return str(out)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    shutil.rmtree(tmp_dir, ignore_errors=True)
    raise RuntimeError(
        "LibreOffice wurde nicht gefunden. "
        "Bitte die .doc-Datei zunächst in Word als .docx speichern."
    )


# ---------------------------------------------------------------------------
# .txt → python-docx Document  (plain-text import)
# ---------------------------------------------------------------------------

def _load_txt_as_document(txt_path: str) -> Document:
    """Read a plain-text file and return a python-docx Document."""
    with open(txt_path, "r", encoding="utf-8", errors="replace") as fh:
        lines = fh.read().splitlines()
    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for line in lines:
        doc.add_paragraph(line)
    return doc


# ---------------------------------------------------------------------------
# Main processing function
# ---------------------------------------------------------------------------

def process_document(
    input_path: str,
    output_path: str,
    track_changes: bool = False,
    ai_engine=None,
    progress_callback=None,
) -> int:
    """Full pipeline: load → detect headings → apply styles → save.

    Returns number of heading paragraphs standardised.
    """

    def progress(msg: str):
        if progress_callback:
            progress_callback(msg)

    # ── Step 1: Prepare input ──────────────────────────────────────────────
    progress("Schritt 1/4 – Datei vorbereiten …")

    # Ensure the output directory exists
    out_dir = Path(output_path).parent
    if not out_dir.exists():
        out_dir.mkdir(parents=True, exist_ok=True)

    path = input_path
    ext = Path(path).suffix.lower()
    _tmp_dir_to_clean: Optional[str] = None

    if ext in (".doc", ".rtf"):
        path = convert_doc_to_docx(path)
        _tmp_dir_to_clean = str(Path(path).parent)
        ext = ".docx"

    # ── Step 2: Load document ──────────────────────────────────────────────
    progress("Schritt 2/4 – Überschriften analysieren …")

    if ext == ".txt":
        doc = _load_txt_as_document(path)
    else:
        try:
            doc = Document(path)
        except Exception as open_err:
            try:
                converted = convert_doc_to_docx(path)
                _tmp_dir_to_clean = str(Path(converted).parent)
                doc = Document(converted)
            except Exception:
                if _tmp_dir_to_clean:
                    shutil.rmtree(_tmp_dir_to_clean, ignore_errors=True)
                raise RuntimeError(
                    f"Die Datei konnte nicht geöffnet werden: {open_err}\n\n"
                    "Mögliche Ursachen:\n"
                    "  • Die .docx-Datei ist im alten .doc-Format gespeichert\n"
                    "  • Die Datei ist passwortgeschützt\n"
                    "  • Die Datei ist beschädigt\n"
                    "  • Die Datei ist noch in Word geöffnet (Sperre)\n\n"
                    "Lösung: Datei in Word öffnen und als .docx neu speichern."
                ) from open_err

    # ── Detect headings ────────────────────────────────────────────────────
    if ai_engine and ai_engine.api_key:
        headings = ai_engine.analyze_headings(doc)
        style_headings = detect_headings(doc)
        for idx, lvl in style_headings.items():
            if idx not in headings:
                headings[idx] = lvl
    else:
        headings = detect_headings(doc)

    headings = normalize_levels(headings)

    # ── Step 3 (no OOXML auto-numbering — literal numbers are written instead)
    progress("Schritt 3/4 – Nummerierung einrichten …")

    # ── Step 4: Apply heading styles + literal 1./1.1/1.1.1 numbers ───────
    progress("Schritt 4/4 – Gliederung standardisieren …")
    apply_heading_styles(doc, headings, track_changes=track_changes)

    doc.save(output_path)

    # Clean up any LibreOffice temp directory created during conversion
    if _tmp_dir_to_clean:
        shutil.rmtree(_tmp_dir_to_clean, ignore_errors=True)

    return len(headings)
